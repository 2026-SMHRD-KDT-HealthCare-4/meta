from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict

class DOCXGenerator:
    def generate(self, job_name: str, edited_sections: Dict[str, str]) -> BytesIO:
        doc = Document()
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(f"자기소개서 - {job_name}")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph() # Spacer
        
        # Sections
        section_titles = {
            "motivation": "1. 지원 동기",
            "competency": "2. 직무 역량",
            "growth": "3. 성장 과정",
            "plan": "4. 입사 후 포부"
        }
        
        for section_key, content in edited_sections.items():
            title_text = section_titles.get(section_key, section_key.capitalize())
            
            p_title = doc.add_paragraph()
            p_title_run = p_title.add_run(title_text)
            p_title_run.bold = True
            p_title_run.font.size = Pt(14)
            
            p_content = doc.add_paragraph(content)
            p_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_paragraph() # Spacer
            
        # Save to BytesIO
        target = BytesIO()
        doc.save(target)
        target.seek(0)
        return target
